# app.py
# AI Legal Assistant — Main Streamlit Application
# Production build — fully local, no external calls

import json
import os
import time
import logging
import tempfile
from pathlib import Path

os.environ["TOKENIZERS_PARALLELISM"] = "false"
logging.getLogger("sentence_transformers").setLevel(logging.ERROR)

import streamlit as st
from dotenv import load_dotenv
from modules.ingestion import CHROMA_PATH, COLLECTION_NAME, EMBEDDING_MODEL, chroma_client
from modules.case_manager import (
    list_cases, get_active_case, set_active_case, create_case, delete_case
)
from modules.feedback import (
    log_feedback, log_redaction_feedback,
    check_and_log_auto_failures, get_feedback_stats
)
from modules.ingestion import ingest_document, get_ingested_documents, clear_all_documents
from modules.redaction import redact_document
from modules.search import search_case_law, lookup_citation  # ⚠️ Requires internet — disabled for Shenelle's deployment
from modules.llm import (
    rag_query, stream_rag_query, summarize_documents,
    get_primary_model,
)
from modules.agentic_rag import stream_agentic_rag_query
from modules.hardware_detect import get_current_profile
from modules.setup_wizard import run_health_check, HealthStatus
from modules.retrieval import normalize_score
from modules.cache import clear_cache

load_dotenv()

# ── Page Config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Legal AI Assistant",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)


# ── Custom CSS ────────────────────────────────────────────────────────────────

st.markdown("""
<style>
    /* ── Base ──────────────────────────────────────────────── */
    html, body, [class*="css"] { font-size: 18px !important; }
    .main, .stApp { background-color: #0D1117; color: #E6EDF3; }

    /* ── Header strip — deep teal ──────────────────────────── */
    header[data-testid="stHeader"] {
        background-color: #0A3D3D !important;
        border-bottom: 1px solid #1F5F5F !important;
    }

    /* ── Sidebar ────────────────────────────────────────────── */
    [data-testid="stSidebar"] { background-color: #5A86AD !important; }
    [data-testid="stSidebar"] * { color: #2C2C2C !important; background-color: transparent !important; }
    [data-testid="stSidebar"] .stMarkdown p { color: #2C2C2C !important; font-size: 18px !important; }
    [data-testid="stSidebar"] code {
        color: #1A1A1A !important;
        background-color: rgba(0,0,0,0.15) !important;
        white-space: nowrap !important;
        overflow: hidden !important;
        text-overflow: ellipsis !important;
        max-width: 200px !important;
        display: inline-block !important;
    }
    [data-testid="stSidebar"] h2 { font-size: 26px !important; color: #1A1A1A !important; }
    [data-testid="stSidebar"] h2:first-of-type { font-size: 32px !important; }
    [data-testid="stSidebar"] h3 { font-size: 20px !important; color: #1A1A1A !important; }

    /* ── Sidebar buttons — transparent, charcoal border ────── */
    [data-testid="stSidebar"] .stButton > button {
        background-color: transparent !important;
        color: #2C2C2C !important;
        border: 1.5px solid #2C2C2C !important;
        border-radius: 6px !important;
        font-size: 17px !important;
    }
    [data-testid="stSidebar"] .stButton > button:hover {
        background-color: rgba(0,0,0,0.1) !important;
        color: #1A1A1A !important;
        border-color: #1A1A1A !important;
    }
    
    /* ── Sidebar selectbox ──────────────────────────────────── */
    [data-testid="stSidebar"] [data-baseweb="select"] > div {
        background-color: transparent !important;
        border: 1.5px solid #2C2C2C !important;
        border-radius: 6px !important;
    }
    
    [data-testid="stSidebar"] [data-baseweb="select"] * {
        color: #2C2C2C !important;
        background-color: transparent !important;
    }

    /* ── Sidebar slider ─────────────────────────────────────── */
    [data-testid="stSidebar"] [data-testid="stSlider"] > div,
    [data-testid="stSidebar"] [data-testid="stSlider"] > div > div,
    [data-testid="stSidebar"] [data-testid="stSlider"] > div > div > div {
        background-color: transparent !important;
    }
    [data-testid="stSidebar"] [data-baseweb="slider"] > div {
        background-color: #3D6B8A !important;
        height: 4px !important;
    }
    [data-testid="stSidebar"] [data-baseweb="slider"] > div > div {
        background-color: #2EA043 !important;
    }
    [data-testid="stSidebar"] [role="slider"] {
        background-color: #2EA043 !important;
        border: 3px solid #2EA043 !important;
        box-shadow: none !important;
    }
    [data-testid="stSidebar"] [data-testid="stSlider"] span { color: #2C2C2C !important; }

    [data-testid="stSidebar"] [role="slider"] {
        background-color: #2EA043 !important;
        border-color: #2EA043 !important;
        box-shadow: none !important;
    }

    /* ── Main title ─────────────────────────────────────────── */
    .stApp h1 { font-size: 42px !important; color: #E6EDF3 !important; }

    /* ── Main area buttons — green ──────────────────────────── */
    section[data-testid="stMain"] .stButton > button {
        background-color: #238636 !important;
        color: #FFFFFF !important;
        border: 2px solid #2EA043 !important;
        font-size: 18px !important;
        border-radius: 6px !important;
        transition: all 0.2s ease !important;
    }
    section[data-testid="stMain"] .stButton > button:hover {
        background-color: #FFFFFF !important;
        color: #238636 !important;
        border-color: #238636 !important;
    }

    /* ── Browse Files button ────────────────────────────────── */
    [data-testid="stFileUploadDropzone"] button,
    [data-testid="stBaseButton-secondary"] {
        background-color: transparent !important;
        color: #2C2C2C !important;
        border: 1.5px solid #2C2C2C !important;
        border-radius: 6px !important;
        font-size: 18px !important;
    }
    [data-testid="stFileUploadDropzone"] button:hover,
    [data-testid="stBaseButton-secondary"]:hover {
        background-color: rgba(0,0,0,0.1) !important;
        color: #1A1A1A !important;
        border-color: #1A1A1A !important;
    }

    /* ── Tabs ───────────────────────────────────────────────── */
    .stTabs [data-baseweb="tab"] {
        font-size: 18px !important;
        color: #E6EDF3 !important;
        background-color: transparent !important;
        border: none !important;
    }
    .stTabs [data-baseweb="tab"]:hover { color: #58A6FF !important; }
    .stTabs [aria-selected="true"] {
        color: #2EA043 !important;
        border-bottom: 3px solid #2EA043 !important;
        background-color: transparent !important;
    }
    .stTabs [data-baseweb="tab-highlight"] { background-color: #2EA043 !important; }
    .stTabs [data-baseweb="tab-border"] { background-color: #2EA043 !important; }

    .stTabs [data-baseweb="tab"] span {
        color: #E6EDF3 !important;
    }
    .stTabs [aria-selected="true"] span {
        color: #2EA043 !important;
    }
    .stTabs [data-baseweb="tab"]:hover span {
        color: #58A6FF !important;
    }

    /* ── Chat messages ──────────────────────────────────────── */
    [data-testid="stChatMessage"] .stMarkdown p,
    [data-testid="stChatMessage"] .stMarkdown li,
    [data-testid="stChatMessage"] .stMarkdown ol li,
    [data-testid="stChatMessage"] .stMarkdown ul li,
    [data-testid="stChatMessage"] .stMarkdown h1,
    [data-testid="stChatMessage"] .stMarkdown h2,
    [data-testid="stChatMessage"] .stMarkdown h3,
    [data-testid="stChatMessage"] .stMarkdown strong,
    [data-testid="stChatMessage"] .stMarkdown em,
    [data-testid="stChatMessage"] .stMarkdown span {
        font-size: 23px !important;
        color: #D0E8FF !important;
        line-height: 1.7 !important;
    }
    [data-testid="stChatMessage"] ol li,
    [data-testid="stChatMessage"] ul li { color: #D0E8FF !important; font-size: 23px !important; }

    /* ── Citation & answer boxes ────────────────────────────── */
    .citation-box {
        background-color: #161B22;
        border-left: 3px solid #238636;
        padding: 10px 15px;
        border-radius: 4px;
        margin: 5px 0;
        font-size: 16px !important;
        font-family: monospace;
        color: #A8D8A8 !important;
    }
    .answer-box {
        background-color: #161B22;
        border: 1px solid #30363D;
        border-radius: 8px;
        padding: 20px;
        margin: 10px 0;
        font-size: 23px !important;
        color: #D0E8FF !important;
        line-height: 1.7 !important;
    }

    /* ── Chat input ─────────────────────────────────────────── */
    .stChatInput input { font-size: 18px !important; }

    /* ── Status labels ──────────────────────────────────────── */
    .status-success { color: #2EA043 !important; font-weight: bold !important; font-size: 18px !important; }
    .status-skip    { color: #5A3E00 !important; font-weight: bold; }
    .status-error   { color: #7A0000 !important; font-weight: bold; }
    [data-testid="stSidebar"] #ollama-status,
    [data-testid="stSidebar"] #ollama-status * {
        color: #2EA043 !important;
        background-color: transparent !important;
    }

    /* ── Redaction checkboxes ───────────────────────────────── */
    [data-testid="stCheckbox"] label,
    [data-testid="stCheckbox"] p { color: #D0E8FF !important; font-size: 18px !important; }

    /* ── Download buttons ───────────────────────────────────── */
    [data-testid="stDownloadButton"] button {
        background-color: transparent !important;
        color: #D0E8FF !important;
        border: 1.5px solid #D0E8FF !important;
        font-size: 18px !important;
        border-radius: 6px !important;
    }
    [data-testid="stDownloadButton"] button:hover {
        background-color: rgba(208,232,255,0.1) !important;
        color: #FFFFFF !important;
        border-color: #FFFFFF !important;
    }


        /* ── Sidebar app title ── */
    [data-testid="stSidebarHeader"] a,
    [data-testid="stSidebarHeader"] span,
    [data-testid="stSidebarHeader"] p,
    .st-emotion-cache-1cypcdb,
    header[data-testid="stHeader"] ~ div [data-testid="stSidebar"] a {
        color: #02C39A !important;
    }

    /* ── Document store file names ── */
    [data-testid="stFileUploaderFile"] *,
    [data-testid="stFileUploaderFile"] small,
    [data-testid="stFileUploaderFile"] span {
        color: #c8ecec !important;
        font-family: 'Share Tech Mono', monospace !important;
        font-size: 11px !important;
    }

    /* ── Slider caption text ── */
    [data-testid="stCaptionContainer"] p,
    [data-testid="stCaptionContainer"] span,
    .stSlider + div p,
    .stSlider ~ [data-testid="stText"] {
        color: #8ababa !important;
        font-size: 11px !important;
    }

        /* ── Hide heading anchor link on hover ── */
    h1 a[href], h2 a[href], h3 a[href] {
        display: none !important;
    }
    [data-testid="stHeadingWithActionElements"] a {
        display: none !important;
    }

    /* ── Hide sidebar collapse button icon text ── */
    [data-testid="stSidebarCollapseButton"] span,
    [data-testid="stSidebarCollapseButton"] div {
        font-size: 0 !important;
        color: transparent !important;
    }

        /* ── Sidebar title 15% larger ── */
    [data-testid="stSidebarHeader"] *,
    [data-testid="stSidebarNav"] * {
        font-size: 1.95rem !important;
    }

        /* ── Sidebar title glow ── */
    [data-testid="stSidebar"] a,
    [data-testid="stSidebar"] a *,
    [data-testid="stSidebar"] a span,
    [data-testid="stSidebarHeader"],
    [data-testid="stSidebarHeader"] * {
        text-shadow: 0 0 12px rgba(255, 255, 255, 0.54),
                    0 0 30px rgba(220, 240, 240, 0.36) !important;
    }

    /* ── Sidebar section labels glow ── */
    [data-testid="stSidebar"] h3 {
        text-shadow: 0 0 8px rgba(255, 255, 255, 0.48),
                    0 0 20px rgba(220, 240, 240, 0.30) !important;
        letter-spacing: 3px !important;
    }

    /* ── System Status <details> collapsible ─────────────────── */
    [data-testid="stSidebar"] details.ss-details {
        border: 1.5px solid #2C2C2C;
        border-radius: 6px;
        overflow: hidden;
        transition: box-shadow 0.2s ease;
    }
    [data-testid="stSidebar"] details.ss-details:hover {
        box-shadow: 0 0 14px rgba(2, 195, 154, 0.30),
                    0 4px 16px rgba(2, 195, 154, 0.12);
    }
    [data-testid="stSidebar"] summary.ss-summary {
        padding: 7px 12px;
        cursor: pointer;
        display: flex !important;
        justify-content: space-between;
        align-items: center;
        list-style: none;
        background: transparent;
    }
    [data-testid="stSidebar"] summary.ss-summary::-webkit-details-marker { display: none; }
    [data-testid="stSidebar"] summary.ss-summary::marker               { display: none; }
    [data-testid="stSidebar"] summary.ss-summary::after {
        content: "▾";
        color: #2C2C2C !important;
        font-size: 0.78rem;
        flex-shrink: 0;
        margin-left: 4px;
    }
    [data-testid="stSidebar"] details.ss-details[open] summary.ss-summary::after { content: "▲"; }
    [data-testid="stSidebar"] .ss-title { font-size: 0.88rem; color: #2C2C2C !important; flex: 1; }
    [data-testid="stSidebar"] .ss-info  { color: #8B949E !important; cursor: help; font-size: 0.8rem; margin-right: 6px; }
    [data-testid="stSidebar"] .ss-body  { border-top: 1.5px solid #2C2C2C; padding: 7px 12px; font-size: 0.75rem; }
    [data-testid="stSidebar"] .ss-model    { color: #03e8b5 !important; font-weight: 600; }
    [data-testid="stSidebar"] .ss-dot-ok  { color: #2EA043 !important; }
    [data-testid="stSidebar"] .ss-dot-err { color: #CF222E !important; }

</style>
""", unsafe_allow_html=True)


# ── Health Check (non-blocking) ───────────────────────────────────────────────
# Runs on every page load. Result is used by the sidebar status card and the
# main-area warning banner. Never calls st.stop() — the app always loads.
_health = run_health_check()


# ── Session State ─────────────────────────────────────────────────────────────

if "messages" not in st.session_state:
    st.session_state.messages = []
if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0
if "show_new_case_form" not in st.session_state:
    st.session_state.show_new_case_form = False
if "confirm_delete" not in st.session_state:
    st.session_state.confirm_delete = False
# Resolve active case on first load
if "active_case" not in st.session_state:
    _cases_init = list_cases()   # triggers legacy auto-register
    _active_init = get_active_case()
    if _active_init is None and _cases_init:
        _active_init = _cases_init[0]["case_id"]
        set_active_case(_active_init)
    st.session_state.active_case = _active_init

if "ingested_docs" not in st.session_state:
    st.session_state.ingested_docs = get_ingested_documents(
        case_id=st.session_state.get("active_case")
    )
# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## ⚖️ Amicus Ai")
    st.markdown("---")

    # ── Case Management ───────────────────────────────────────────────────────
    st.markdown("### Active Case")
    _all_cases = list_cases()

    if not _all_cases:
        st.markdown("*No cases yet.*")
        if st.button("+ Create First Case", key="create_first_btn"):
            st.session_state.show_new_case_form = True
    else:
        _case_ids    = [c["case_id"] for c in _all_cases]
        _case_labels = {c["case_id"]: c["display_name"] for c in _all_cases}
        _active      = st.session_state.get("active_case") or _case_ids[0]
        if _active not in _case_ids:
            _active = _case_ids[0]
            st.session_state.active_case = _active
            set_active_case(_active)

        _active_meta = next(c for c in _all_cases if c["case_id"] == _active)
        st.markdown(
            f'<div style="background:rgba(2,195,154,0.15);border-left:3px solid #02C39A;'
            f'padding:8px 12px;border-radius:4px;margin-bottom:8px;">'
            f'<strong style="color:#02C39A !important;">{_active_meta["display_name"]}</strong>'
            f'</div>',
            unsafe_allow_html=True,
        )

        if len(_all_cases) > 1:
            _selected = st.selectbox(
                "Switch Case",
                options=_case_ids,
                format_func=lambda x: _case_labels[x],
                index=_case_ids.index(_active),
                key="case_selector",
            )
            if _selected != _active:
                st.session_state.active_case = _selected
                set_active_case(_selected)
                st.session_state.messages = []
                st.session_state.ingested_docs = []
                st.rerun()

        _cb1, _cb2 = st.columns(2)
        with _cb1:
            if st.button("+ New Case", key="new_case_btn"):
                st.session_state.show_new_case_form = not st.session_state.show_new_case_form
                st.session_state.confirm_delete = False
        with _cb2:
            if st.button("🗑️ Delete", key="delete_case_btn"):
                st.session_state.confirm_delete = not st.session_state.confirm_delete
                st.session_state.show_new_case_form = False

    # New Case form
    if st.session_state.get("show_new_case_form"):
        with st.form("new_case_form", clear_on_submit=True):
            _nc_id   = st.text_input("Case ID (lowercase, underscores)", placeholder="e.g. smith_v_acme")
            _nc_name = st.text_input("Display Name", placeholder="e.g. Smith v. Acme Corp")
            _nc_desc = st.text_area("Description (optional)", height=60)
            if st.form_submit_button("Create Case"):
                if not _nc_id or not _nc_name:
                    st.error("Case ID and Display Name are required.")
                else:
                    try:
                        create_case(_nc_id.strip(), _nc_name.strip(), _nc_desc.strip())
                        st.session_state.active_case = _nc_id.strip()
                        st.session_state.show_new_case_form = False
                        st.session_state.messages = []
                        st.session_state.ingested_docs = []
                        st.rerun()
                    except ValueError as _e:
                        st.error(str(_e))

    # Delete confirmation
    if st.session_state.get("confirm_delete") and _all_cases:
        _del_name = next(
            (c["display_name"] for c in _all_cases if c["case_id"] == _active), _active
        )
        st.warning(f"Delete **{_del_name}**? All documents will be permanently removed.")
        _dc1, _dc2 = st.columns(2)
        with _dc1:
            if st.button("Yes, Delete", key="confirm_del_yes"):
                delete_case(_active)
                st.session_state.active_case = get_active_case()
                st.session_state.confirm_delete = False
                st.session_state.messages = []
                st.session_state.ingested_docs = []
                st.rerun()
        with _dc2:
            if st.button("Cancel", key="confirm_del_no"):
                st.session_state.confirm_delete = False
                st.rerun()

    st.markdown("---")

    # Document upload — only available when a case is active
    _sidebar_case_id = st.session_state.get("active_case")
    st.markdown("### Upload Documents")
    if not _sidebar_case_id:
        st.markdown("*Create a case above to upload documents.*")
    else:
        uploaded_files = st.file_uploader(
            "Drop files here",
            type=["pdf", "docx", "xlsx", "xls", "txt", "csv"],
            accept_multiple_files=True,
            help="Supported: PDF, Word, Excel, TXT, CSV",
            key=f"uploader_{st.session_state.uploader_key}"
        )

        if uploaded_files:
            for uploaded_file in uploaded_files:
                if uploaded_file.name not in st.session_state.ingested_docs:
                    with st.spinner(f"Processing {uploaded_file.name}..."):
                        with tempfile.NamedTemporaryFile(
                            delete=False,
                            suffix=Path(uploaded_file.name).suffix
                        ) as tmp:
                            tmp.write(uploaded_file.getbuffer())
                            tmp_path = tmp.name
                        try:
                            result = ingest_document(
                                tmp_path,
                                original_name=uploaded_file.name,
                                case_id=_sidebar_case_id,
                            )
                        finally:
                            os.unlink(tmp_path)

                        if result["status"] == "success":
                            st.markdown(
                                f'<p class="status-success">✓ {uploaded_file.name}'
                                f' ({result["chunks"]} chunks)</p>',
                                unsafe_allow_html=True
                            )
                            st.session_state.ingested_docs = get_ingested_documents(
                                case_id=_sidebar_case_id
                            )
                        else:
                            st.markdown(
                                f'<p class="status-skip">⚠ {uploaded_file.name}'
                                f' — {result["status"]}</p>',
                                unsafe_allow_html=True
                            )

    st.markdown("---")

    # Document inventory
    st.markdown("### Documents in Store")
    _docs_all = get_ingested_documents(case_id=_sidebar_case_id)
    _active_case_display = next(
        (c["display_name"] for c in _all_cases if c["case_id"] == _sidebar_case_id),
        _sidebar_case_id or "—",
    )

    def _doc_pill(name: str) -> str:
        label = (name[:22] + "…") if len(name) > 25 else name
        return (
            f'<div style="background:rgba(2,195,154,0.10);border:1px solid rgba(2,195,154,0.28);'
            f'border-radius:10px;padding:3px 10px;margin:2px 0;font-size:0.76rem;'
            f'font-family:monospace;color:#02C39A;white-space:nowrap;overflow:hidden;'
            f'text-overflow:ellipsis;">{label}</div>'
        )

    if _docs_all:
        _n = len(_docs_all)
        st.markdown(
            f'<div style="font-size:0.78rem;color:#5A86AD;margin-bottom:6px;">'
            f'{_n} document{"s" if _n != 1 else ""}&nbsp;·&nbsp;'
            f'Active case: <strong>{_active_case_display}</strong>'
            f'</div>',
            unsafe_allow_html=True,
        )

        for _d in _docs_all[:4]:
            st.markdown(_doc_pill(_d), unsafe_allow_html=True)

        if _n > 4:
            if "show_all_docs" not in st.session_state:
                st.session_state.show_all_docs = False
            if st.button(
                f"View all {_n} documents ↓", key="view_all_docs_btn", use_container_width=True
            ):
                st.session_state.show_all_docs = not st.session_state.show_all_docs
            if st.session_state.get("show_all_docs"):
                _doc_filter = st.text_input(
                    "", placeholder="Search documents...",
                    key="doc_search", label_visibility="collapsed",
                )
                _filtered_docs = (
                    [d for d in _docs_all if _doc_filter.lower() in d.lower()]
                    if _doc_filter else _docs_all
                )
                for _d in _filtered_docs:
                    st.markdown(_doc_pill(_d), unsafe_allow_html=True)

        st.markdown("---")
        if st.button("🗑️ Clear All Documents", key="clear_docs"):
            clear_all_documents(case_id=_sidebar_case_id)
            clear_cache()
            st.session_state.ingested_docs = []
            st.session_state.messages = []
            st.session_state.uploader_key += 1
            st.rerun()
    else:
        st.markdown("*No documents uploaded yet*")

    # Settings
    st.markdown("### Settings")
    top_k = st.slider("Speed vs. Performance", min_value=3, max_value=10, value=7)
    st.caption("Lower provides a faster response — higher provides a more in-depth response at the cost of speed.")
    _am_lbl_col, _am_ico_col = st.columns([0.85, 0.15])
    with _am_lbl_col:
        st.markdown("**Analysis Mode**")
    with _am_ico_col:
        st.markdown(" ", help="""\
- Basic — Fast, reliable analysis. Best for straightforward document lookup and single-file queries.

- Advanced — Highest precision. Retrieves and reranks the most relevant passages before answering. Recommended for most cases.

- Expert — Deep cross-document reasoning. Best for complex queries that require connecting facts across multiple files. Adds ~20s processing time.""")
    _mode_label = st.selectbox(
        "Analysis Mode",
        options=["Basic", "Advanced", "Expert"],
        index=1,
        label_visibility="collapsed",
        key="retrieval_mode_label"
    )
    _mode_map = {"Basic": "hybrid", "Advanced": "rerank", "Expert": "agentic"}
    retrieval_mode = _mode_map[_mode_label]

    # ── System Status — IT health indicator (collapsed by default) ────────────
    st.markdown("<div style='margin-top:20px'></div>", unsafe_allow_html=True)
    _MODEL_DISPLAY_NAMES = {
        "llama3.1:8b":      "Llama 3.1 8B",
        "llama3.3:8b":      "Llama 3.3 8B",
        "mistral-nemo:12b": "Mistral Nemo 12B",
        "llama3.1:70b":     "Llama 3.1 70B",
    }
    _ss_model         = get_primary_model()
    _ss_model_display = _MODEL_DISPLAY_NAMES.get(_ss_model, _ss_model or "Standard")
    _ss_dot_class     = "ss-dot-ok" if _health.ollama_running else "ss-dot-err"

    # Row 1 — label left, ? popover right (mirrors Analysis Mode pattern)
    _ss_lbl_col, _ss_ico_col = st.columns([0.85, 0.15])
    with _ss_lbl_col:
        st.markdown("**System Status**")
    with _ss_ico_col:
        st.markdown(" ", help=(
            "System Status shows whether Amicus is ready to work.\n\n"
            "● Green dot — Everything is running normally. "
            "Amicus is ready to analyze your documents.\n\n"
            "● Red dot — Amicus is having trouble starting up. "
            "Try closing and reopening the app. "
            "If the problem continues, contact your IT administrator."
        ))

    # Row 2 — bordered collapsible box with chevron inside
    st.markdown(
        f'<details class="ss-details">'
        f'<summary class="ss-summary">'
        f'<span class="ss-title">System Status</span>'
        f'</summary>'
        f'<div class="ss-body">'
        f'<span class="ss-model">{_ss_model_display}</span>'
        f'&nbsp;&nbsp;|&nbsp;&nbsp;'
        f'<span class="{_ss_dot_class}">●</span>'
        f'</div>'
        f'</details>',
        unsafe_allow_html=True,
    )


# ── Main Area ─────────────────────────────────────────────────────────────────
# ── Visual Polish & Sparkles ──────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@600;700&family=Share+Tech+Mono&display=swap');

/* ── Scale everything up 20% ── */
html { font-size: 19.2px !important; }

/* ── Gradient title ── */
h1 {
    font-family: 'Rajdhani', sans-serif !important;
    font-weight: 700 !important;
    font-size: 2.6rem !important;
    background: linear-gradient(135deg, #ffffff 0%, #03e8b5 50%, #02a8b8 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    letter-spacing: 0.03em !important;
    filter: drop-shadow(0 0 18px rgba(2, 195, 154, 0.35));
}

/* ── Subheadings ── */
h2 {
    font-family: 'Rajdhani', sans-serif !important;
    font-weight: 700 !important;
    color: #03e8b5 !important;
    letter-spacing: 0.05em !important;
    font-size: 1.6rem !important;
}
h3 {
    font-family: 'Share Tech Mono', monospace !important;
    font-size: 0.8rem !important;
    letter-spacing: 3px !important;
    text-transform: uppercase !important;
    color: #02C39A !important;
}

/* ── Sidebar title ── */
[data-testid="stSidebarHeader"] *,
[data-testid="stSidebarNav"] * {
    color: #03e8b5 !important;
    font-family: 'Rajdhani', sans-serif !important;
    font-weight: 700 !important;
    font-size: 1.3rem !important;
}

/* ── Sidebar accent line ── */
[data-testid="stSidebar"]::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 2px;
    background: linear-gradient(90deg, transparent, #02C39A, #028090, transparent);
    z-index: 999;
}

/* ── Active tab glow ── */
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {
    color: #03e8b5 !important;
    text-shadow: 0 0 12px rgba(2, 195, 154, 0.8) !important;
    border-bottom: 2px solid #02C39A !important;
}

/* ── Query input glow on focus ── */
[data-testid="stTextArea"] textarea:focus,
[data-testid="stTextInput"] input:focus {
    box-shadow: 0 0 0 1px rgba(2, 195, 154, 0.5),
                0 0 25px rgba(2, 195, 154, 0.15) !important;
    border-color: #02C39A !important;
}

/* ── Primary button shimmer ── */
[data-testid="stButton"] > button[kind="primary"] {
    font-family: 'Rajdhani', sans-serif !important;
    font-weight: 700 !important;
    letter-spacing: 0.08em !important;
    text-transform: uppercase !important;
    transition: all 0.25s ease !important;
    box-shadow: 0 4px 15px rgba(2, 195, 154, 0.3) !important;
}
[data-testid="stButton"] > button[kind="primary"]:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 28px rgba(2, 195, 154, 0.5) !important;
}

/* ── Citation expanders — left accent bar ── */
[data-testid="stExpander"] {
    border-left: 3px solid #02C39A !important;
    border-radius: 0 10px 10px 0 !important;
    transition: all 0.25s ease !important;
}
[data-testid="stExpander"]:hover {
    box-shadow: -3px 0 18px rgba(2, 195, 154, 0.25),
                0 4px 20px rgba(2, 195, 154, 0.1) !important;
    transform: translateX(2px) !important;
}

/* ── Slider thumb glow ── */
[data-testid="stSlider"] [role="slider"] {
    box-shadow: 0 0 10px rgba(2, 195, 154, 0.8),
                0 0 25px rgba(2, 195, 154, 0.4) !important;
}

/* ── Document file name tags ── */
[data-testid="stFileUploaderFile"] * {
    font-family: 'Share Tech Mono', monospace !important;
    color: #03e8b5 !important;
    font-size: 0.75rem !important;
}

/* ── Custom scrollbar ── */
::-webkit-scrollbar { width: 5px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb {
    background: linear-gradient(180deg, #02C39A, #028090);
    border-radius: 3px;
}

/* ── Fade-up entrance on content ── */
[data-testid="stVerticalBlock"] > div {
    animation: fadeUp 0.35s ease-out both;
}
@keyframes fadeUp {
    from { opacity: 0; transform: translateY(6px); }
    to   { opacity: 1; transform: translateY(0); }
}
</style>

<script>
(function() {
    const canvas = document.createElement('canvas');
    canvas.id = 'amicus-canvas';
    canvas.style.cssText = 'position:fixed;top:0;left:0;width:100vw;height:100vh;pointer-events:none;z-index:9999;opacity:0.5;';
    document.body.appendChild(canvas);
    const ctx = canvas.getContext('2d');
    function resize() { canvas.width = window.innerWidth; canvas.height = window.innerHeight; }
    resize();
    window.addEventListener('resize', resize);

    const COLORS = ['#02C39A','#028090','#00A896','#F4C430','#ffffff','#5DCAA5'];
    const particles = [];

    class P {
        constructor(x, y, burst) {
            this.x = x !== undefined ? x : Math.random() * canvas.width;
            this.y = y !== undefined ? y : canvas.height + 10;
            this.size = Math.random() * (burst ? 3 : 2) + 0.5;
            this.vy = -(Math.random() * (burst ? 2.5 : 0.5) + (burst ? 0.5 : 0.15));
            this.vx = (Math.random() - 0.5) * (burst ? 3 : 0.3);
            this.alpha = Math.random() * 0.7 + 0.2;
            this.color = COLORS[Math.floor(Math.random() * COLORS.length)];
            this.twinkle = Math.random() * Math.PI * 2;
            this.ts = Math.random() * 0.03 + 0.01;
            this.star = Math.random() > 0.65;
            this.burst = burst || false;
            this.life = burst ? 120 : Infinity;
        }
        update() {
            this.x += this.vx;
            this.y += this.vy;
            this.twinkle += this.ts;
            this.currentAlpha = this.alpha * (0.5 + 0.5 * Math.sin(this.twinkle));
            if (this.burst) { this.life--; this.currentAlpha *= (this.life / 120); }
            return this.y > -20 && (!this.burst || this.life > 0);
        }
        draw() {
            ctx.save();
            ctx.globalAlpha = Math.max(0, this.currentAlpha);
            ctx.fillStyle = this.color;
            ctx.shadowBlur = 6;
            ctx.shadowColor = this.color;
            if (this.star) {
                ctx.translate(this.x, this.y);
                ctx.beginPath();
                for (let i = 0; i < 5; i++) {
                    const a1 = (i * 72 - 90) * Math.PI / 180;
                    const a2 = (i * 72 - 54) * Math.PI / 180;
                    ctx.lineTo(Math.cos(a1) * this.size * 1.6, Math.sin(a1) * this.size * 1.6);
                    ctx.lineTo(Math.cos(a2) * this.size * 0.6, Math.sin(a2) * this.size * 0.6);
                }
                ctx.closePath();
                ctx.fill();
            } else {
                ctx.beginPath();
                ctx.arc(this.x, this.y, this.size, 0, Math.PI * 2);
                ctx.fill();
            }
            ctx.restore();
        }
    }

    for (let i = 0; i < 70; i++) {
        const p = new P();
        p.y = Math.random() * canvas.height;
        particles.push(p);
    }

    function loop() {
        ctx.clearRect(0, 0, canvas.width, canvas.height);
        for (let i = particles.length - 1; i >= 0; i--) {
            if (!particles[i].update()) {
                if (!particles[i].burst) particles[i] = new P();
                else particles.splice(i, 1);
            } else {
                particles[i].draw();
            }
        }
        requestAnimationFrame(loop);
    }
    loop();

    document.addEventListener('click', e => {
        for (let i = 0; i < 14; i++) particles.push(new P(e.clientX, e.clientY, true));
    });
})();
</script>
""", unsafe_allow_html=True)

st.markdown("# ⚖️ Document Analysis")
st.markdown("*Fully local · Air-gapped · Attorney-client privilege protected*")

if not _health.ollama_running:
    st.warning(
        "⚠️ Analysis engine is not responding. "
        "Please ensure Ollama is running before submitting queries. "
        "See `install/QUICK_START.md` for setup instructions."
    )

# ── Onboarding — shown when no cases exist ────────────────────────────────────
if not list_cases() or not st.session_state.get("active_case"):
    st.markdown("---")
    st.markdown(
        """
        <div style="text-align:center; padding: 60px 20px;">
            <div style="font-size: 72px; margin-bottom: 24px;">⚖️</div>
            <h2 style="color:#02C39A; font-size:2rem; margin-bottom:16px;">
                Welcome to Amicus AI
            </h2>
            <p style="font-size:1.2rem; color:#A0AEC0; max-width:540px; margin:0 auto 32px;">
                Amicus AI is a fully local, air-gapped legal research assistant.
                Each matter lives in its own isolated document store — nothing crosses case boundaries.
            </p>
            <p style="font-size:1.1rem; color:#E6EDF3;">
                👈 Use the <strong>Active Case</strong> panel in the sidebar to create your first case.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.stop()

tab1, tab2, tab3, tab4 = st.tabs(["💬 Query Documents", "📋 Summarize", "🛡️ Redact", "🔍 Case Law"])

# ── TAB 1: Query ──────────────────────────────────────────────────────────────
with tab1:
    st.markdown("### Ask a Question About Your Documents")

    # Chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message.get("sources"):
                st.markdown("**Sources:**")
                for source in message["sources"]:
                    st.markdown(
                        f'<div class="citation-box">'
                        f'📄 {source["file"]} — Page {source["page"]} '
                        f'(confidence: {normalize_score(source)}/10)'
                        f'</div>',
                        unsafe_allow_html=True
                    )

    # Clear conversation button
    if st.session_state.get("messages"):
        if st.button("🗑️ Clear Conversation", key="clear_chat"):
            st.session_state.messages = []
            st.rerun()

    # Query input
    if query := st.chat_input("e.g. Who are the parties in this case?"):
        docs = get_ingested_documents(case_id=st.session_state.get("active_case"))
        if not docs:
            st.warning("Please upload documents first.")
        else:
            # Add user message
            st.session_state.messages.append({"role": "user", "content": query})
            with st.chat_message("user"):
                st.markdown(query)

            # Get RAG response — streaming
            _mode        = retrieval_mode
            _top_k       = top_k
            _case_id     = st.session_state.get("active_case")
            result       = {}
            _query_start = time.time()

            _mode_label_map = {
                "hybrid": "Basic", "rerank": "Advanced", "agentic": "Expert"
            }

            with st.chat_message("assistant"):
                # Two containers fix render order: status on top, answer below
                _status_area = st.container()   # pipeline stages (top)
                _answer_area = st.container()   # streaming answer (bottom)

                with _answer_area:
                    _answer_ph = st.empty()

                if _mode == "agentic":
                    # Expert mode: LangGraph sub-stage progress via st.status()
                    with _status_area:
                        with st.status(
                            "🧠 Planning sub-queries...", expanded=False
                        ) as _agentic_status:
                            def _agentic_progress(msg: str) -> None:
                                _agentic_status.update(label=msg, state="running")

                            _streamed_text = ""
                            for _tok in stream_agentic_rag_query(
                                query,
                                top_k=_top_k,
                                case_id=_case_id,
                                progress_callback=_agentic_progress,
                                result_holder=result,
                            ):
                                _streamed_text += _tok
                                _answer_ph.markdown(_streamed_text + "▌")

                            _agentic_status.update(label="✅ Complete", state="complete")

                    _final = result.get("answer", _streamed_text)
                    _answer_ph.markdown(_final)

                    if _final.strip() != _streamed_text.strip():
                        st.info("Response revised after quality review.")

                else:
                    # Basic / Advanced: retrieve → (rank) → generate with stage labels
                    with _status_area:
                        with st.status(
                            "🔍 Searching documents...", expanded=False
                        ) as _status:
                            _t0  = time.time()
                            gen  = stream_rag_query(
                                query,
                                top_k=_top_k,
                                mode=_mode,
                                case_id=_case_id,
                                result_holder=result,
                            )

                            # Block until first token — retrieval (+ rerank) runs here
                            _first_tok = next(gen, None)

                            # Ensure "searching" stage is visible ≥ 0.3 s
                            _rl = time.time() - _t0
                            if _rl < 0.3:
                                time.sleep(0.3 - _rl)

                            if _mode == "rerank":
                                _status.update(
                                    label="📊 Ranking evidence...", state="running"
                                )
                                time.sleep(0.3)

                            _status.update(
                                label="✍️ Generating response...", state="running"
                            )

                            _full_answer = _first_tok or ""
                            if _full_answer:
                                _answer_ph.markdown(_full_answer + "▌")

                            for _tok in gen:
                                _full_answer += _tok
                                _answer_ph.markdown(_full_answer + "▌")

                            _status.update(label="✅ Complete", state="complete")

                    _final = result.get("answer", _full_answer)
                    _answer_ph.markdown(_final)

                # ── Metadata row: time · chunks · mode · case ─────────────────
                _elapsed_total = time.time() - _query_start
                _case_display  = next(
                    (c["display_name"] for c in _all_cases if c["case_id"] == _case_id),
                    _case_id or "—",
                )
                st.markdown(
                    f'<div style="margin-top:8px;padding:6px 12px;'
                    f'background:rgba(255,255,255,0.04);border:1px solid #30363D;'
                    f'border-radius:6px;color:#8B949E;font-size:0.78rem;'
                    f'font-family:monospace;">'
                    f'⏱ {_elapsed_total:.1f}s &nbsp;·&nbsp; '
                    f'📄 {result.get("chunks_used", 0)} chunks &nbsp;·&nbsp; '
                    f'🔧 {_mode_label_map.get(_mode, _mode)} &nbsp;·&nbsp; '
                    f'📁 {_case_display}'
                    f'</div>',
                    unsafe_allow_html=True,
                )

                if result.get("sources"):
                    st.markdown("**Sources:**")
                    for source in result["sources"]:
                        st.markdown(
                            f'<div class="citation-box">'
                            f'📄 {source["file"]} — Page {source["page"]} '
                            f'(confidence: {normalize_score(source)}/10)'
                            f'</div>',
                            unsafe_allow_html=True,
                        )

                # Citation verification panel
                _cr = result.get("citation_report")
                if _cr is not None:
                    _score = _cr.overall_confidence_score
                    if _score >= 0.8:
                        _badge_color, _badge_label = "#238636", "High"
                    elif _score >= 0.5:
                        _badge_color, _badge_label = "#9e6a03", "Medium"
                    else:
                        _badge_color, _badge_label = "#da3633", "Low"

                    st.markdown(
                        f'<div style="margin-top:10px;padding:10px 14px;'
                        f'background:rgba(255,255,255,0.03);border:1px solid #30363D;'
                        f'border-radius:6px;font-size:0.82rem;">'
                        f'<span style="font-weight:600;color:#C9D1D9;">Citation Verification</span>'
                        f'&nbsp;&nbsp;'
                        f'<span style="background:{_badge_color};color:#fff;padding:2px 8px;'
                        f'border-radius:4px;font-size:0.75rem;font-weight:600;">'
                        f'{_badge_label} confidence ({_score:.0%})</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

                    for cit in _cr.verified_citations:
                        st.markdown(
                            f'<div class="citation-box" style="border-left:3px solid #238636;">'
                            f'✅ {cit["raw"]}</div>',
                            unsafe_allow_html=True,
                        )

                    for cit in _cr.unverified_citations:
                        st.markdown(
                            f'<div class="citation-box" style="border-left:3px solid #da3633;">'
                            f'⚠️ {cit["raw"]} — '
                            f'<em>Citation not found in retrieved documents — verify manually</em>'
                            f'</div>',
                            unsafe_allow_html=True,
                        )

                    if _cr.missing_citations:
                        with st.expander(f"⚠️ {len(_cr.missing_citations)} unsupported factual claim(s) detected"):
                            for claim in _cr.missing_citations:
                                st.markdown(f"- `{claim}`")

            # Auto-failure detection
            _conf_scores = [
                normalize_score(s)
                for s in result.get("sources", [])
            ]

            _doc_list = [s["file"] for s in result.get("sources", [])]
            check_and_log_auto_failures(
                query=query,
                response=result.get("answer", ""),
                mode=_mode,
                top_k=st.session_state.get("top_k", 7),
                chunks_used=result.get("chunks_used", 0),
                confidence_scores=_conf_scores,
                document_list=_doc_list,
                tab="query",
            )

            # Store result in session state for feedback UI persistence
            st.session_state["last_result"] = {
                "query": query,
                "answer": result.get("answer", ""),
                "sources": result.get("sources", []),
                "chunks_used": result.get("chunks_used", 0),
                "mode": _mode,
                "top_k": st.session_state.get("top_k", 7),
                "conf_scores": _conf_scores,
                "doc_list": _doc_list,
            }
            st.session_state["show_feedback"] = True
            st.session_state["feedback_submitted"] = False

            # Save to history
            st.session_state.messages.append({
                "role": "assistant",
                "content": result.get("answer", ""),
                "sources": result.get("sources", []),
            })

    # ── Feedback UI — rendered outside chat block, persists on rerun ──────────
    if st.session_state.get("show_feedback") and not st.session_state.get("feedback_submitted"):
        lr = st.session_state.get("last_result", {})
        if lr:
            st.markdown("---")
            st.caption("Was this response helpful?")
            fb_col1, fb_col2, _ = st.columns([1, 1, 8])
            with fb_col1:
                if st.button("👍", key="fb_up", help="Good response"):
                    log_feedback(
                        feedback_type="thumbs_up",
                        query=lr["query"],
                        response=lr["answer"],
                        mode=lr["mode"],
                        top_k=lr["top_k"],
                        sources=lr["sources"],
                        chunks_used=lr["chunks_used"],
                        confidence_scores=lr["conf_scores"],
                        document_list=lr["doc_list"],
                        tab="query",
                    )
                    st.session_state["show_feedback"] = False
                    st.session_state["feedback_submitted"] = True
                    st.toast("Thanks for the feedback! ✔")
                    st.rerun()
            with fb_col2:
                if st.button("👎", key="fb_down", help="Poor response"):
                    st.session_state["show_thumbs_down_comment"] = True

            if st.session_state.get("show_thumbs_down_comment"):
                _comment = st.text_input(
                    "What was wrong? (optional)",
                    key="fb_comment",
                    placeholder="e.g. missed the termination date, wrong party name..."
                )
                if st.button("Submit feedback", key="fb_submit"):
                    log_feedback(
                        feedback_type="thumbs_down",
                        query=lr["query"],
                        response=lr["answer"],
                        mode=lr["mode"],
                        top_k=lr["top_k"],
                        sources=lr["sources"],
                        chunks_used=lr["chunks_used"],
                        confidence_scores=lr["conf_scores"],
                        document_list=lr["doc_list"],
                        comment=_comment if _comment else None,
                        tab="query",
                    )
                    st.session_state["show_feedback"] = False
                    st.session_state["feedback_submitted"] = True
                    st.session_state["show_thumbs_down_comment"] = False
                    st.toast("Feedback recorded. Thank you. ✔")
                    st.rerun()

# ── TAB 2: Summarize ──────────────────────────────────────────────────────────
with tab2:
    st.markdown("### Generate Document Summary")
    st.markdown("Produces a structured executive summary of all uploaded documents.")

    docs = get_ingested_documents(case_id=st.session_state.get("active_case"))
    if not docs:
        st.warning("Please upload documents first.")
    else:
        st.markdown("**Documents to summarize:**")
        for doc in docs:
            st.markdown(f"📄 `{doc}`")

        if st.button("Generate Summary", type="primary"):
            with st.spinner("Analyzing documents... this may take 30-60 seconds."):
                summary = summarize_documents()

            st.markdown("### Summary")
            st.markdown(
                f'<div class="answer-box">{summary}</div>',
                unsafe_allow_html=True
            )

            # Download summary as TXT
            st.download_button(
                label="Download Summary as TXT",
                data=summary,
                file_name="case_summary.txt",
                mime="text/plain"
            )
# ── TAB 3: Redact ─────────────────────────────────────────────────────────────
with tab3:
    st.markdown("### 🛡️ Document Redaction & Template Generator")
    st.markdown("Upload a document to automatically detect and replace PII with placeholders.")

    all_categories = {
        "[PARTY_NAME]":       "Person names",
        "[ORGANIZATION]":     "Companies & courts",
        "[LOCATION]":         "Cities & states",
        "[DATE]":             "All dates",
        "[CASE_NO]":          "Case numbers",
        "[SSN]":              "Social security numbers",
        "[PHONE]":            "Phone numbers",
        "[EMAIL]":            "Email addresses",
        "[FINANCIAL_AMOUNT]": "Dollar amounts",
        "[BAR_NO]":           "Attorney bar numbers",
        "[ZIP_CODE]":         "ZIP codes",
        "[DOCKET_NO]":        "Docket numbers",
    }

    col1, col2 = st.columns([2, 1])
    
    with col2:
        st.markdown("**Select Categories to Redact**")
        st.caption("Nothing is redacted unless explicitly selected.")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("☑️ Select All", key="select_all",
                        use_container_width=True):
                for placeholder in all_categories:
                    st.session_state[f"redact_{placeholder}"] = True
        with col_b:
            if st.button("☐ Clear All", key="clear_all_redact",
                        use_container_width=True):
                for placeholder in all_categories:
                    st.session_state[f"redact_{placeholder}"] = False

        selected_categories = []
        for placeholder, description in all_categories.items():
            if st.checkbox(
                f"{placeholder} — {description}",
                value=False,
                key=f"redact_{placeholder}"
            ):
                selected_categories.append(placeholder)

        st.markdown("---")
        aggressive = st.toggle("Aggressive Mode", value=False)
        st.caption("Also redacts legal references and cardinal numbers.")
    
    with col1:
        redact_file = st.file_uploader(
            "Upload document to redact",
            type=["pdf", "docx", "xlsx", "txt", "csv"],
            key="redact_uploader",
            help="Document will be processed locally — no data leaves this machine"
        )

        if redact_file:
            if not selected_categories:
                st.info("☝️ Select at least one category above to enable redaction.")

            if st.button("🛡️ Redact Document", type="primary",
                        disabled=len(selected_categories) == 0):
                with st.spinner(f"Redacting {redact_file.name}..."):
                    # Save to temp file
                    with tempfile.NamedTemporaryFile(
                        delete=False,
                        suffix=Path(redact_file.name).suffix
                    ) as tmp:
                        tmp.write(redact_file.getbuffer())
                        tmp_path = tmp.name

                    # Redact
                    output_path = tmp_path + "_REDACTED.txt"
                    report = redact_document(
                        tmp_path,
                        output_path=output_path,
                        aggressive=aggressive,
                        categories=selected_categories if selected_categories else None
                    )
                    os.unlink(tmp_path)

                if report["status"] == "success":
                    st.success(f"✅ Redaction complete — {report['total_redactions']} items redacted")
                # if report["status"] == "success":
                #     st.success(f"✅ Redaction complete — {report['total_redactions']} items redacted | categories={selected_categories} | log_len={len(report.get('redactions', []))}")
                    
                    # Show placeholder counts
                    st.markdown("**Redaction Summary:**")
                    cols = st.columns(3)
                    for i, (placeholder, count) in enumerate(report['placeholder_counts'].items()):
                        with cols[i % 3]:
                            st.metric(placeholder, count)

                    # Read and show preview
                    with open(output_path, "r", encoding="utf-8") as f:
                        redacted_content = f.read()

                    st.markdown("**Preview (first 1000 characters):**")
                    st.markdown(
                        f'<div class="answer-box" style="font-size:16px !important;">'
                        f'{redacted_content[:1000]}...'
                        f'</div>',
                        unsafe_allow_html=True
                    )

                    # Export options
                    st.markdown("**Export Redacted Document As:**")
                    export_col1, export_col2, export_col3 = st.columns(3)

                    with export_col1:
                        st.download_button(
                            label="📄 Export as TXT",
                            data=redacted_content,
                            file_name=f"{Path(redact_file.name).stem}_REDACTED.txt",
                            mime="text/plain",
                            use_container_width=True
                        )

                    with export_col2:
                        from docx import Document as DocxDocument
                        from docx.shared import Pt, RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from io import BytesIO
                        doc = DocxDocument()

                        # Header
                        title = doc.add_heading("REDACTED DOCUMENT", level=1)
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        meta = doc.add_paragraph()
                        meta.add_run(f"Source: ").bold = True
                        meta.add_run(redact_file.name)
                        meta2 = doc.add_paragraph()
                        meta2.add_run(f"Total redactions: ").bold = True
                        meta2.add_run(str(report['total_redactions']))
                        doc.add_paragraph("─" * 60)

                        # Body with highlighted placeholders
                        import re
                        placeholder_pattern = re.compile(r'\[([A-Z_]+)\]')
                        for line in redacted_content.split("\n"):
                            if not line.strip():
                                continue
                            para = doc.add_paragraph()
                            last_end = 0
                            for match in placeholder_pattern.finditer(line):
                                # Text before placeholder
                                if match.start() > last_end:
                                    para.add_run(line[last_end:match.start()])
                                # Highlighted placeholder
                                run = para.add_run(match.group())
                                run.bold = True
                                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                                run.font.highlight_color = 7  # yellow
                                last_end = match.end()
                            # Remaining text
                            if last_end < len(line):
                                para.add_run(line[last_end:])

                        # Placeholder legend at end
                        doc.add_paragraph("─" * 60)
                        legend_title = doc.add_paragraph()
                        legend_title.add_run("PLACEHOLDER LEGEND").bold = True
                        legend = {
                            "[PARTY_NAME]": "Person / party name",
                            "[ORGANIZATION]": "Company, court, or institution",
                            "[LOCATION]": "City, state, or address",
                            "[DATE]": "Date or date range",
                            "[CASE_NO]": "Case or docket number",
                            "[SSN]": "Social security number",
                            "[PHONE]": "Phone number",
                            "[EMAIL]": "Email address",
                            "[FINANCIAL_AMOUNT]": "Dollar amount",
                            "[BAR_NO]": "Attorney bar number",
                        }
                        for placeholder, description in legend.items():
                            if placeholder in redacted_content:
                                p = doc.add_paragraph(style='List Bullet')
                                r = p.add_run(placeholder)
                                r.bold = True
                                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                                p.add_run(f" — {description}")

                        docx_buffer = BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        st.download_button(
                            label="📝 Export as DOCX Template",
                            data=docx_buffer.getvalue(),
                            file_name=f"{Path(redact_file.name).stem}_TEMPLATE.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )

                    with export_col3:
                        # Generate PDF in memory
                        from reportlab.lib.pagesizes import letter
                        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                        from reportlab.lib.styles import getSampleStyleSheet
                        from io import BytesIO as BytesIO2
                        pdf_buffer = BytesIO2()
                        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                        styles = getSampleStyleSheet()
                        story = []
                        story.append(Paragraph("REDACTED DOCUMENT", styles['Title']))
                        story.append(Paragraph(f"Source: {redact_file.name}", styles['Normal']))
                        story.append(Paragraph(f"Total redactions: {report['total_redactions']}", styles['Normal']))
                        story.append(Spacer(1, 12))
                        for para in redacted_content.split("\n"):
                            if para.strip():
                                try:
                                    story.append(Paragraph(para, styles['Normal']))
                                    story.append(Spacer(1, 6))
                                except Exception:
                                    pass
                        pdf_doc.build(story)
                        pdf_buffer.seek(0)
                        st.download_button(
                            label="📋 Export as PDF",
                            data=pdf_buffer.getvalue(),
                            file_name=f"{Path(redact_file.name).stem}_REDACTED.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )

                    # Store redaction result for persistent feedback UI
                    st.session_state["last_redact"] = {
                        "document_name": redact_file.name,
                        "categories_selected": selected_categories,
                        "redaction_count": report["total_redactions"],
                    }
                    st.session_state["show_redact_feedback"] = True
                    st.session_state["redact_feedback_submitted"] = False
                    os.unlink(output_path)
                else:
                    st.error(f"Redaction failed: {report['status']}")

        # ── Redact feedback UI — outside result block, persists on rerun ─────────
    if st.session_state.get("show_redact_feedback") and not st.session_state.get("redact_feedback_submitted"):
        lr = st.session_state.get("last_redact", {})
        if lr:
            st.markdown("---")
            st.caption("Was the redaction accurate?")
            red_col1, red_col2, _ = st.columns([1, 1, 8])
            with red_col1:
                if st.button("👍", key="redact_up", help="Redaction looks correct"):
                    log_redaction_feedback(
                        feedback_type="thumbs_up",
                        document_name=lr["document_name"],
                        categories_selected=lr["categories_selected"],
                        redaction_count=lr["redaction_count"],
                    )
                    st.session_state["show_redact_feedback"] = False
                    st.session_state["redact_feedback_submitted"] = True
                    st.toast("Thanks for the feedback! ✔")
                    st.rerun()
            with red_col2:
                if st.button("👎", key="redact_down", help="Redaction missed something"):
                    st.session_state["show_redact_comment"] = True
            if st.session_state.get("show_redact_comment"):
                _red_comment = st.text_input(
                    "What was missed or incorrectly redacted? (optional)",
                    key="redact_comment",
                    placeholder="e.g. missed SSN on page 3, kept attorney name..."
                )
                if st.button("Submit", key="redact_submit"):
                    log_redaction_feedback(
                        feedback_type="thumbs_down",
                        document_name=lr["document_name"],
                        categories_selected=lr["categories_selected"],
                        redaction_count=lr["redaction_count"],
                        comment=_red_comment if _red_comment else None,
                    )
                    st.session_state["show_redact_feedback"] = False
                    st.session_state["redact_feedback_submitted"] = True
                    st.session_state["show_redact_comment"] = False
                    st.toast("Feedback recorded. Thank you. ✔")
                    st.rerun()

# ── TAB 4: Case Law Search ────────────────────────────────────────────────────
with tab4:
    st.markdown("### 🔍 Case Law Search")
    if "search_enabled" not in st.session_state:
        st.session_state.search_enabled = False
    col_toggle1, col_toggle2 = st.columns([1, 5])
    with col_toggle1:
        if st.session_state.search_enabled:
            if st.button("🔴 Deactivate", type="secondary"):
                st.session_state.search_enabled = False
                st.rerun()
        else:
            if st.button("🟢 Activate", type="primary"):
                st.session_state.search_enabled = True
                st.rerun()
    with col_toggle2:
        if st.session_state.search_enabled:
            st.error("🔓 **Air-gap is OFF** — Network active. Do not upload client documents during this session.", icon="⚠️")
        else:
            st.warning("🔒 **Activating will disable air-gap security.** Clear all documents from the store before proceeding.", icon="🛡️")
    SEARCH_ENABLED = st.session_state.search_enabled

    col1, col2 = st.columns([3, 1])

    with col1:
        case_query = st.text_input(
            "Search case law",
            placeholder="e.g. ADA reasonable accommodation interactive process Second Circuit",
            disabled=not SEARCH_ENABLED
        )

    with col2:
        court_filter = st.selectbox(
            "Filter by court",
            options=["All courts", "scotus", "ca1", "ca2", "ca3", "ca4",
                     "ca5", "ca6", "ca7", "ca8", "ca9", "ca10", "ca11"],
            disabled=not SEARCH_ENABLED
        )
        max_results = st.slider("Results", 3, 10, 5, disabled=not SEARCH_ENABLED)

    if st.button("🔍 Search Case Law", type="primary", disabled=not SEARCH_ENABLED):
        if not case_query:
            st.warning("Enter a search query.")
        else:
            with st.spinner("Searching CourtListener..."):
                from modules.search import search_case_law
                court = None if court_filter == "All courts" else court_filter
                results = search_case_law(case_query, court=court, max_results=max_results)

            if results and "error" in results[0]:
                st.error(results[0]["error"])
            else:
                st.success(f"Found {len(results)} results")
                for r in results:
                    with st.expander(f"📄 {r['case_name']} — {r['citation']}"):
                        col_a, col_b = st.columns(2)
                        with col_a:
                            st.markdown(f"**Court:** {r['court']}")
                            st.markdown(f"**Date:** {r['date_filed']}")
                        with col_b:
                            st.markdown(f"**Status:** {r.get('status', 'N/A')}")
                            st.markdown(f"[View on CourtListener]({r['url']})")
                        if r.get("summary"):
                            st.markdown("**Excerpt:**")
                            st.markdown(
                                f'<div class="citation-box">{r["summary"]}</div>',
                                unsafe_allow_html=True
                            )

    # Citation lookup
    st.markdown("---")
    st.markdown("**Look up a specific citation:**")
    cite_col1, cite_col2 = st.columns([3, 1])
    with cite_col1:
        citation_input = st.text_input(
            "Citation",
            placeholder="e.g. 737 F.3d 834",
            disabled=not SEARCH_ENABLED,
            key="citation_lookup"
        )
    with cite_col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Look Up", disabled=not SEARCH_ENABLED):
            if citation_input:
                with st.spinner("Looking up citation..."):
                    from modules.search import lookup_citation
                    result = lookup_citation(citation_input)
                if "error" in result:
                    st.error(result["error"])
                else:
                    st.success(f"**{result['case_name']}** — {result['citation']}")
                    st.markdown(f"Court: {result['court']} | Date: {result['date_filed']}")
                    st.markdown(f"[View full opinion]({result['url']})")