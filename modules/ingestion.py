# modules/ingestion.py
# Document upload, text extraction, chunking, and ChromaDB storage pipeline
import os
import logging
os.environ["TOKENIZERS_PARALLELISM"] = "false"
logging.getLogger("sentence_transformers").setLevel(logging.ERROR)

import os
import hashlib
from pathlib import Path
from typing import List, Dict, Any

import chromadb
import pymupdf
from docx import Document as DocxDocument
from openpyxl import load_workbook
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer

# ── Configuration ─────────────────────────────────────────────────────────────
CHROMA_PATH = "./db/chroma"
COLLECTION_NAME = "legal_docs"
CHUNK_SIZE = 800
CHUNK_OVERLAP = 300
EMBEDDING_MODEL = "all-MiniLM-L6-v2"  # fast, local, no external calls

# ── Document Type Detection ────────────────────────────────────────────────────
def detect_document_type(source: str, text_sample: str) -> str:
    """Infer document type from filename and content for metadata enrichment."""
    source_lower = source.lower()
    text_lower = text_sample[:500].lower()

    if "deposition" in source_lower or "depo" in source_lower:
        return "deposition"
    if "complaint" in source_lower:
        return "complaint"
    if "answer" in source_lower and "affirmative" in text_lower:
        return "answer"
    if "email" in source_lower or "email_chain" in source_lower:
        return "email_chain"
    if "pip" in source_lower or "performance_improvement" in source_lower:
        return "pip"
    if "accommodation" in source_lower:
        return "accommodation"
    if "termination" in source_lower:
        return "termination"
    if "eeoc" in source_lower or "charge" in source_lower:
        return "eeoc_charge"
    if "medical" in source_lower or "okonkwo" in source_lower:
        return "medical_record"
    if "timeline" in source_lower:
        return "timeline"
    if "witness" in source_lower or "statement" in source_lower:
        return "witness_statement"
    if "damages" in source_lower or "wages" in source_lower:
        return "damages"
    if "deposition" in text_lower and ("q:" in text_lower or "a:" in text_lower):
        return "deposition"
    if "email" in text_lower and ("from:" in text_lower or "subject:" in text_lower):
        return "email_chain"
    return "legal_document"

# ── Initialize Components ─────────────────────────────────────────────────────
embedding_model = SentenceTransformer(EMBEDDING_MODEL, cache_folder="./db/embeddings")
chroma_client = chromadb.PersistentClient(path=CHROMA_PATH)

def _get_collection():
    """Single source of truth for the ChromaDB collection."""
    return chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"}
    )

def clean_ocr_text(text: str) -> str:
    """
    Fix common Tesseract OCR artifacts in legal documents.
    Two-pass approach:
    Pass 1: Regex fixes for structural artifacts (fast, deterministic)
    Pass 2: LLM cleanup for residual errors (slower, only if OCR flag is set)
    """
    import re

    # ── Pass 1: Regex fixes ───────────────────────────────────────────────────

    # Fix missing space between lowercase and uppercase (todismiss → to dismiss)
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)

    # Fix missing space after period before capital (order.The → order. The)
    text = re.sub(r'([a-z])\.([A-Z])', r'\1. \2', text)

    # Fix missing space after comma (Smith,Jones → Smith, Jones)
    text = re.sub(r'([a-zA-Z]),([a-zA-Z])', r'\1, \2', text)

    # Fix run-together court names (DISTRICTCOURT → DISTRICT COURT)
    common_joins = [
        (r'DISTRICT(COURT|OF)', r'DISTRICT \1'),
        (r'SOUTHERN(DISTRICT)', r'SOUTHERN \1'),
        (r'NORTHERN(DISTRICT)', r'NORTHERN \1'),
        (r'EASTERN(DISTRICT)', r'EASTERN \1'),
        (r'WESTERN(DISTRICT)', r'WESTERN \1'),
        (r'UNITED(STATES)', r'UNITED \1'),
        (r'CIRCUIT(COURT)', r'CIRCUIT \1'),
        (r'(COURT)(OF)', r'\1 \2'),
        (r'([A-Z]{2,})([A-Z][a-z]{3,})', r'\1 \2'),  # TORRESDistrict → TORRES District
        (r'MOTION(TO|FOR|IN)', r'MOTION \1'),
        (r'(?i)(pursuant)(to|ly)', r'\1 \2')
    ]
    for pattern, replacement in common_joins:
        text = re.sub(pattern, replacement, text)

    # Fix missing space before legal keywords run together
    legal_keywords = [
        "pursuant", "plaintiff", "defendant", "motion", "order",
        "whereas", "hereby", "thereof", "therein", "dismiss", "state", 
        "grant", "relief", "claim", "failure", "proceed", "allege", 
        "contend", "argue", "assert", "pursuant", "relief", "can", 
        "judge", "district", "wherein"
    ]
   
    for kw in legal_keywords:
        text = re.sub(rf'([a-z])({kw})', rf'\1 \2', text, flags=re.IGNORECASE)

    # Normalize multiple spaces and blank lines
    text = re.sub(r' {2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()

def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract text from PDF with OCR fallback for scanned pages.
    Strategy:
    1. Try native text extraction (fast, perfect for digital PDFs)
    2. If page has no extractable text, run Tesseract OCR via PyMuPDF
    """
    pages = []
    doc = pymupdf.open(file_path)
    
    for page_num, page in enumerate(doc, start=1):
        # Stage 1: Try native text extraction
        text = page.get_text().strip()
        
        if text and len(text) > 20:
            # Digital PDF page — use native text
            pages.append({
                "text": text,
                "page": page_num,
                "ocr": False
            })
        else:
            # Stage 2: Scanned page — run OCR via Tesseract
            try:
                tp = page.get_textpage_ocr(flags=0, language="eng", dpi=300, full=True)
                ocr_text = page.get_text(textpage=tp).strip()
                if ocr_text and len(ocr_text) > 20:
                    ocr_text = clean_ocr_text(ocr_text)
                    pages.append({
                        "text": ocr_text,
                        "page": page_num,
                        "ocr": True
                    })
                    logging.info(f"OCR extracted {len(ocr_text)} chars from page {page_num}")
                else:
                    logging.warning(f"OCR returned no text for page {page_num}")
            except Exception as e:
                logging.error(f"OCR failed on page {page_num}: {e}")
    
    doc.close()
    return pages


def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from Word documents including tables and headers."""
    doc = DocxDocument(file_path)
    text_parts = []

    # Standard paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                text_parts.append(row_text)

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        return []
    return [{"text": full_text, "page": 1}]


def extract_text_from_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from Excel spreadsheets, sheet by sheet."""
    pages = []
    wb = load_workbook(file_path, read_only=True, data_only=True)
    for sheet_num, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) for cell in row if cell is not None])
            if row_text.strip():
                rows.append(row_text)
        if rows:
            pages.append({
                "text": f"[Sheet: {sheet_name}]\n" + "\n".join(rows),
                "page": sheet_num
            })
    wb.close()
    return pages


def extract_text_from_txt(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from plain text files."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    return [{"text": text, "page": 1}]

def extract_text_from_csv(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from CSV files."""
    import csv
    rows = []
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = " | ".join([cell for cell in row if cell.strip()])
            if row_text:
                rows.append(row_text)
    if not rows:
        return []
    return [{"text": "\n".join(rows), "page": 1}]



def extract_text(file_path: str) -> List[Dict[str, Any]]:
    """Route to correct extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    extractors = {
        ".pdf":  extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".xls":  extract_text_from_xlsx,
        ".txt":  extract_text_from_txt,
        ".csv":  extract_text_from_csv,
    }
    if ext not in extractors:
        raise ValueError(f"Unsupported file type: {ext}")
    return extractors[ext](file_path)


def chunk_pages(pages: List[Dict], source: str) -> List[Dict[str, Any]]:
    """
    Split extracted pages into chunks with metadata.
    Uses document-type-aware chunking:
    - Depositions: smaller chunks (400 chars) preserving Q&A turns with speaker context
    - Email chains: split on email boundaries preserving sender/recipient
    - Other legal docs: standard chunking with section-aware separators
    """
    # Detect document type from filename and content sample
    sample_text = pages[0]["text"] if pages else ""
    doc_type = detect_document_type(source, sample_text)

    # Deposition-specific separators — preserve Q&A turn boundaries
    deposition_separators = [
        "\nQ:", "\nA:",          # Q&A turns — primary split point
        "\n\nQ:", "\n\nA:",      # Double newline before Q&A
        "\n\n\n",
        "\n\n",
        "\n", ". ", " ", ""
    ]

    # Email chain separators — preserve individual email boundaries
    email_separators = [
        "\n────────",            # Our email divider
        "\nFROM:",               # Email header
        "\nDate:",
        "\n\n\n",
        "\n\n",
        "\n", ". ", " ", ""
    ]

    # Standard legal document separators
    legal_separators = [
        "\n\n\n",
        "\n\n",
        "\nWHEREAS", "\nNOW, THEREFORE", "\nIN WITNESS",
        "\n1. ", "\n2. ", "\n3. ", "\n4. ", "\n5. ",
        "\n6. ", "\n7. ", "\n8. ", "\n9. ", "\n10. ",
        "\nI. ", "\nII. ", "\nIII. ", "\nIV. ", "\nV. ",
        "\nA. ", "\nB. ", "\nC. ", "\nD. ", "\nE. ",
        "\nCOUNT ", "\nCLAIM ", "\nCAUSE OF ACTION",
        "\n\n", "\n", ". ", " ", ""
    ]

    # Choose chunk size and separators based on document type
    if doc_type == "deposition":
        chunk_size = 400      # Smaller — preserve individual Q&A exchanges
        chunk_overlap = 150   # Overlap carries speaker context forward
        separators = deposition_separators
    elif doc_type == "email_chain":
        chunk_size = 600      # Medium — preserve individual email context
        chunk_overlap = 100
        separators = email_separators
    else:
        chunk_size = CHUNK_SIZE
        chunk_overlap = CHUNK_OVERLAP
        separators = legal_separators

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=separators
    )

    chunks = []
    for page in pages:
        splits = splitter.split_text(page["text"])
        for i, split in enumerate(splits):
            if not split.strip():
                continue
            chunks.append({
                "text":          split,
                "source":        source,
                "page":          page["page"],
                "chunk_index":   i,
                "ocr":           page.get("ocr", False),
                "document_type": doc_type,
            })
    return chunks

# ── Embed and Store Chunks ────────────────────────────────────────────────────
def embed_and_store(chunks: List[Dict], doc_id: str) -> int:
    """Embed chunks locally and store in ChromaDB. Returns chunk count."""
    
    # Deduplicate chunks — prevents repeated template text from dominating
    seen = set()
    unique_chunks = []
    for chunk in chunks:
        text_key = chunk["text"].strip()[:200]
        if text_key not in seen:
            seen.add(text_key)
            unique_chunks.append(chunk)
    
    if len(unique_chunks) < len(chunks):
        print(f"  Deduplicated: {len(chunks)} → {len(unique_chunks)} chunks")
    
    chunks = unique_chunks
    
    texts = [c["text"] for c in chunks]
    embeddings = embedding_model.encode(texts, show_progress_bar=False).tolist()

    ids = [f"{doc_id}_chunk_{i}" for i in range(len(chunks))]
    metadatas = [{"source": c["source"], "page": c["page"], "document_type": c.get("document_type", "legal_document")} for c in chunks]

    _get_collection().upsert(
        ids=ids,
        embeddings=embeddings,
        documents=texts,
        metadatas=metadatas
    )
    return len(chunks)

# ── Ingest Document ───────────────────────────────────────────────────────────
def ingest_document(file_path: str, original_name: str = None) -> Dict[str, Any]:
    """
    Full ingestion pipeline: extract → chunk → embed → store.
    Returns a summary dict with ingestion results.
    """
    file_name = original_name if original_name else Path(file_path).name
    doc_id = hashlib.md5(file_name.encode()).hexdigest()[:12]

    print(f"  Extracting text from {file_name}...")
    pages = extract_text(file_path)

    print(f"  Chunking {len(pages)} page(s)...")
    chunks = chunk_pages(pages, source=file_name)

    if not chunks:
        return {
            "file": file_name,
            "pages": len(pages),
            "chunks": 0,
            "status": "skipped — no extractable text (possibly image-based PDF)"
        }

    print(f"  Embedding and storing {len(chunks)} chunks...")
    stored = embed_and_store(chunks, doc_id)

    return {
        "file": file_name,
        "pages": len(pages),
        "chunks": stored,
        "status": "success"
    }


def get_ingested_documents() -> List[str]:
    """Return list of unique source documents currently in the vector store."""
    results = _get_collection().get(include=["metadatas"])
    sources = set(m["source"] for m in results["metadatas"] if m)
    return sorted(list(sources))


def clear_all_documents() -> None:
    """Wipe the entire ChromaDB collection — removes all ingested documents."""
    chroma_client.delete_collection(COLLECTION_NAME)
    _get_collection()  # Recreate fresh empty collection